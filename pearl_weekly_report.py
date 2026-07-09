import os
from datetime import datetime, timedelta, timezone

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

from src.drive.drive import upload_file, download_file_by_name
from src.utils.duplicate import add_duplicate_keys, remove_similar_titles


CAMBODIA_TZ = timezone(timedelta(hours=7))

OUTPUT_WEEKLY = "data/weekly"
OUTPUT_MASTER = "data/master"

MASTER_FILENAME = "PEARL_master_news.csv"
MASTER_FILE = f"{OUTPUT_MASTER}/{MASTER_FILENAME}"


def clean_duplicates(df):
    if df.empty:
        return df

    df = add_duplicate_keys(df)
    df = df.drop_duplicates(subset=["title_id"])
    df = df.drop_duplicates(subset=["url_id"])
    df = remove_similar_titles(df, threshold=0.92)
    return df


def remove_timezone(df):
    for col in df.columns:
        if pd.api.types.is_datetime64tz_dtype(df[col]):
            df[col] = df[col].dt.tz_localize(None)
    return df


def add_weekly_page(doc, title, df, max_items=12):
    doc.add_heading(title, level=1)

    if df.empty:
        doc.add_paragraph("No major news identified.")
        return

    doc.add_paragraph(f"Total unique articles this week: {len(df)}")

    if "crop" in df.columns:
        crop_summary = ", ".join(
            [f"{crop}: {count}" for crop, count in df["crop"].value_counts().items()]
        )
        doc.add_paragraph(f"Crop summary: {crop_summary}")

    if "topic" in df.columns:
        topic_summary = ", ".join(
            [f"{topic}: {count}" for topic, count in df["topic"].value_counts().items()]
        )
        doc.add_paragraph(f"Topic summary: {topic_summary}")

    doc.add_heading("Key News", level=2)

    for _, row in df.head(max_items).iterrows():
        news_title = str(row.get("title", "")).strip()
        summary = str(row.get("Summary", "")).strip()
        crop = str(row.get("crop", "")).strip()
        topic = str(row.get("topic", "")).strip()

        if news_title:
            doc.add_paragraph(f"• {news_title}")

        if summary:
            doc.add_paragraph(f"  {summary}")

        doc.add_paragraph(f"  Crop: {crop} | Topic: {topic}")


def main():
    os.makedirs(OUTPUT_WEEKLY, exist_ok=True)
    os.makedirs(OUTPUT_MASTER, exist_ok=True)

    download_file_by_name(MASTER_FILENAME, MASTER_FILE)

    if not os.path.exists(MASTER_FILE):
        print("No master file found. Run daily collector first.")
        return

    df = pd.read_csv(MASTER_FILE)

    if df.empty:
        print("Master file is empty.")
        return

    today = datetime.now(CAMBODIA_TZ).date()
    start_date = today - timedelta(days=7)
    report_date = today.strftime("%Y-%m-%d")

    if "published_date" not in df.columns:
        print("Missing column: published_date")
        return

    df["published_dt"] = pd.to_datetime(
        df["published_date"],
        errors="coerce",
        utc=True
    )

    df["published_dt_kh"] = df["published_dt"].dt.tz_convert(CAMBODIA_TZ)

    weekly = df[
        df["published_dt_kh"].notna()
        & (df["published_dt_kh"].dt.date >= start_date)
        & (df["published_dt_kh"].dt.date <= today)
    ].copy()

    if weekly.empty:
        print("No weekly records found.")
        return

    weekly = clean_duplicates(weekly)
    weekly = remove_timezone(weekly)

    weekly_xlsx = f"{OUTPUT_WEEKLY}/PEARL_weekly_news_{report_date}.xlsx"
    weekly_docx = f"{OUTPUT_WEEKLY}/PEARL_weekly_summary_{report_date}.docx"

    weekly.to_excel(weekly_xlsx, index=False)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    style = doc.styles["Normal"]
    style.font.size = Pt(8)

    doc.add_heading("PEARL Weekly Agriculture News Summary", 0)
    doc.add_paragraph(f"Report date: {report_date}")
    doc.add_paragraph(f"Coverage period: {start_date} to {today}")
    doc.add_paragraph(
        "Two-page summary. News links are excluded from this Word report and kept in Excel."
    )

    cambodia = weekly[
        weekly["country"].astype(str).str.lower() == "cambodia"
    ].copy()

    global_news = weekly[
        weekly["country"].astype(str).str.lower() != "cambodia"
    ].copy()

    add_weekly_page(doc, "Page 1: Cambodia News", cambodia, max_items=12)
    doc.add_page_break()
    add_weekly_page(doc, "Page 2: Global News", global_news, max_items=12)

    doc.save(weekly_docx)

    upload_file(weekly_xlsx)
    upload_file(weekly_docx)

    print("Weekly report completed successfully.")
    print(f"Weekly Excel: {weekly_xlsx}")
    print(f"Weekly Word: {weekly_docx}")


if __name__ == "__main__":
    main()
