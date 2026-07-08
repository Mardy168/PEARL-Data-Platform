import os
from datetime import datetime, timedelta, timezone

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

from src.drive.drive import upload_file, download_file_by_name
from src.utils.duplicate import add_duplicate_keys, remove_similar_titles


CAMBODIA_TZ = timezone(timedelta(hours=7))

OUTPUT_MONTHLY = "data/monthly"
OUTPUT_MASTER = "data/master"

MASTER_FILENAME = "PEARL_master_news.csv"
MASTER_FILE = f"{OUTPUT_MASTER}/{MASTER_FILENAME}"


def add_monthly_page(doc, title, df, max_items=15):
    doc.add_heading(title, level=1)

    if df.empty:
        doc.add_paragraph("No major news identified.")
        return

    doc.add_paragraph(f"Total unique articles this month: {len(df)}")

    if "crop" in df.columns:
        crop_summary = ", ".join(
            [f"{crop}: {count}" for crop, count in df["crop"].value_counts().items()]
        )
        doc.add_paragraph(f"Crop summary: {crop_summary}")

    if "topic" in df.columns:
        topic_summary = ", ".join(
            [f"{topic}: {count}" for topic, count in df["topic"].value_counts().items()]
        )
        doc.add_paragraph(f"Topic summary: {topic_summary}")

    doc.add_heading("Key News", level=2)

    for _, row in df.head(max_items).iterrows():
        news_title = str(row.get("title", "")).strip()
        summary = str(row.get("Summary", "")).strip()
        crop = str(row.get("crop", "")).strip()
        topic = str(row.get("topic", "")).strip()

        if news_title:
            doc.add_paragraph(f"• {news_title}")

        if summary:
            doc.add_paragraph(f"  {summary}")

        doc.add_paragraph(f"  Crop: {crop} | Topic: {topic}")


def add_statistics_page(doc, df):
    doc.add_heading("Page 3: Overall Statistics", level=1)
    doc.add_paragraph(f"Total unique articles this month: {len(df)}")

    if "country" in df.columns:
        doc.add_heading("By Country", level=2)
        for country, count in df["country"].value_counts().items():
            doc.add_paragraph(f"- {country}: {count}")

    if "crop" in df.columns:
        doc.add_heading("By Crop", level=2)
        for crop, count in df["crop"].value_counts().items():
            doc.add_paragraph(f"- {crop}: {count}")

    if "topic" in df.columns:
        doc.add_heading("By Topic", level=2)
        for topic, count in df["topic"].value_counts().items():
            doc.add_paragraph(f"- {topic}: {count}")


def main():
    os.makedirs(OUTPUT_MONTHLY, exist_ok=True)
    os.makedirs(OUTPUT_MASTER, exist_ok=True)

    download_file_by_name(MASTER_FILENAME, MASTER_FILE)

    if not os.path.exists(MASTER_FILE):
        print("No master file found. Run daily collector first.")
        return

    df = pd.read_csv(MASTER_FILE)

    if df.empty:
        print("Master file is empty.")
        return

    today = datetime.now(CAMBODIA_TZ).date()
    report_month = today.strftime("%Y-%m")

    df["date_collected_dt"] = pd.to_datetime(
        df["date_collected"],
        errors="coerce"
    ).dt.date

    monthly = df[
        pd.to_datetime(df["date_collected"], errors="coerce").dt.strftime("%Y-%m") == report_month
    ].copy()

    if monthly.empty:
        print("No monthly records found.")
        return

    monthly = add_duplicate_keys(monthly)
    monthly = monthly.drop_duplicates(subset=["title_id"])
    monthly = monthly.drop_duplicates(subset=["url_id"])
    monthly = remove_similar_titles(monthly, threshold=0.92)

    monthly_xlsx = f"{OUTPUT_MONTHLY}/PEARL_monthly_news_{report_month}.xlsx"
    monthly_docx = f"{OUTPUT_MONTHLY}/PEARL_monthly_summary_{report_month}.docx"

    monthly.to_excel(monthly_xlsx, index=False)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    style = doc.styles["Normal"]
    style.font.size = Pt(8)

    doc.add_heading("PEARL Monthly Agriculture News Summary", 0)
    doc.add_paragraph(f"Report month: {report_month}")
    doc.add_paragraph(
        "Monthly summary. News links are excluded from this Word report and kept in Excel."
    )

    cambodia = monthly[
        monthly["country"].astype(str).str.lower() == "cambodia"
    ].copy()

    global_news = monthly[
        monthly["country"].astype(str).str.lower() != "cambodia"
    ].copy()

    add_monthly_page(doc, "Page 1: Cambodia Monthly News", cambodia, max_items=15)

    doc.add_page_break()

    add_monthly_page(doc, "Page 2: Global Monthly News", global_news, max_items=15)

    doc.add_page_break()

    add_statistics_page(doc, monthly)

    doc.save(monthly_docx)

    upload_file(monthly_xlsx)
    upload_file(monthly_docx)

    print("Monthly report completed successfully.")
    print(f"Monthly Excel: {monthly_xlsx}")
    print(f"Monthly Word: {monthly_docx}")


if __name__ == "__main__":
    main()
