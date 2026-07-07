from datetime import datetime, timezone
from pathlib import Path
import pandas as pd
from docx import Document
from src.drive.upload import upload_file, list_files, download_file
from src.utils.config import OUTPUT_DIR


def download_recent_daily_csvs(limit=10):
    files = list_files(prefix="PEARL_daily_news_", max_results=50)
    csv_files = [f for f in files if f["name"].endswith(".csv")][:limit]
    local_dir = OUTPUT_DIR / "_downloaded_daily"
    local_dir.mkdir(parents=True, exist_ok=True)
    paths = []
    for f in csv_files:
        local_path = local_dir / f["name"]
        try:
            download_file(f["id"], str(local_path))
            paths.append(local_path)
        except Exception as e:
            print(f"Could not download {f['name']}: {e}")
    return paths


def add_crop_section(doc, df, crop):
    doc.add_heading(crop, level=2)
    crop_df = df[df["crop"] == crop].head(12)
    if crop_df.empty:
        doc.add_paragraph("No relevant records found this week.")
        return
    for _, row in crop_df.iterrows():
        doc.add_paragraph(
            f"- {row.get('title','')} | Source: {row.get('source','')} | "
            f"Group: {row.get('source_group','')} | Topic: {row.get('topic','')}",
            style=None
        )


def run_weekly():
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    local_files = list(Path("output").glob("*/PEARL_daily_news_*.csv"))
    downloaded_files = download_recent_daily_csvs(limit=10)
    all_files = list({str(p): p for p in local_files + downloaded_files}.values())

    frames = []
    for f in all_files:
        try:
            frames.append(pd.read_csv(f))
        except Exception as e:
            print(f"Could not read {f}: {e}")

    if not frames:
        print("No CSV files found. Run the daily workflow first.")
        return

    df = pd.concat(frames, ignore_index=True).drop_duplicates(subset=["url"])
    if "relevance_score" in df.columns:
        df = df.sort_values("relevance_score", ascending=False)

    report_dir = OUTPUT_DIR / "weekly_reports"
    report_dir.mkdir(parents=True, exist_ok=True)
    path = report_dir / f"PEARL_weekly_commonality_report_{today}.docx"

    doc = Document()
    doc.add_heading("PEARL Weekly Commonality Intelligence Report", 0)
    doc.add_paragraph("Commonality crops: Mango, Cashew, Rice and Vegetables")
    doc.add_paragraph("Coverage: Cambodia news and global trend signals from Google News RSS, GDELT and curated RSS sources.")

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        f"This weekly report summarizes {len(df)} unique records collected for PEARL commonality crops. "
        "The records are classified by crop, topic, source group and relevance score."
    )

    doc.add_heading("2. Key Findings by Crop", level=1)
    for crop in ["Mango", "Cashew", "Rice", "Vegetables"]:
        add_crop_section(doc, df, crop)

    doc.add_heading("3. Topic Distribution", level=1)
    if "topic" in df.columns:
        topic_counts = df["topic"].fillna("general").value_counts().head(20)
        for topic, count in topic_counts.items():
            doc.add_paragraph(f"- {topic}: {count} records")

    doc.add_heading("4. References", level=1)
    for i, (_, row) in enumerate(df.head(100).iterrows(), start=1):
        doc.add_paragraph(f"[{i}] {row.get('citation', row.get('url',''))}")

    doc.save(path)
    upload_file(str(path), path.name)
    print(f"Weekly report completed: {path}")
