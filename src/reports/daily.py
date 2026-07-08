from docx import Document
from docx.shared import Inches, Pt


def add_daily_section(doc, title, df, max_items=6):
    doc.add_heading(title, level=1)

    if df.empty:
        doc.add_paragraph("No major news identified.")
        return

    doc.add_paragraph(f"Total articles: {len(df)}")

    for _, row in df.head(max_items).iterrows():
        news_title = str(row.get("title", "")).strip()
        summary = str(row.get("Summary", "")).strip()
        crop = str(row.get("crop", "")).strip()
        topic = str(row.get("topic", "")).strip()

        if news_title:
            doc.add_paragraph(f"• {news_title}")

        if summary:
            doc.add_paragraph(f"  {summary}")

        doc.add_paragraph(f"  Crop: {crop} | Topic: {topic}")


def create_daily_word_report(df, output_path, report_date):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    style = doc.styles["Normal"]
    style.font.size = Pt(8)

    doc.add_heading("PEARL Daily Agriculture News Summary", 0)
    doc.add_paragraph(f"Report date: {report_date}")
    doc.add_paragraph(
        "One-page daily summary. News links are excluded from this Word report and kept in Excel."
    )

    cambodia = df[df["country"].astype(str).str.lower() == "cambodia"].copy()
    global_news = df[df["country"].astype(str).str.lower() != "cambodia"].copy()

    add_daily_section(doc, "Cambodia News", cambodia, max_items=5)
    add_daily_section(doc, "Global News", global_news, max_items=5)

    doc.save(output_path)
    return output_path
