from __future__ import annotations

from docx import Document
from docx.shared import Inches, Pt
import pandas as pd


def configure_document(title: str, subtitle: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.4); sec.bottom_margin = Inches(0.4)
    sec.left_margin = Inches(0.5); sec.right_margin = Inches(0.5)
    doc.styles["Normal"].font.size = Pt(8)
    doc.add_heading(title, 0)
    doc.add_paragraph(subtitle)
    doc.add_paragraph("News links are intentionally excluded from this Word report and remain in Excel.")
    return doc


def add_news_section(doc: Document, heading: str, df: pd.DataFrame, max_items: int) -> None:
    doc.add_heading(heading, level=1)
    if df.empty:
        doc.add_paragraph("No major news identified for this section.")
        return
    doc.add_paragraph(f"Total unique articles: {len(df)}")
    for _, row in df.head(max_items).iterrows():
        title = str(row.get("title", "")).strip()
        summary = str(row.get("Summary", "")).strip()
        meta = " | ".join(x for x in [str(row.get("crop", "")).strip(),
                                           str(row.get("topic", "")).strip(),
                                           str(row.get("publisher_name", "")).strip()] if x)
        if title:
            doc.add_paragraph(f"• {title}")
        if summary:
            doc.add_paragraph(f"  {summary}")
        if meta:
            doc.add_paragraph(f"  {meta}")
